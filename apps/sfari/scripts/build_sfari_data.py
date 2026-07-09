"""Build the three SFARI data JSONs from ``docs/SFARI_Clean.docx``.

Emits into ``sfari/data/``:
  * ``sfari-outcome-mapping.json`` — 20 functions -> {physical,chemical,biological} D/i/-
  * ``sfari-functions.json``       — 20 functions (category, order, name, definition, statement)
  * ``sfari-metrics.json``         — ~82 metrics (functionId, name, scale, likertCriteria[], desktop binding)

This is a one-time DEV build step (not run at request time). Use the EASI venv
python, which already has python-docx 1.2.0::

    D:/Code/Work/easi_claude/.venv/Scripts/python.exe sfari/scripts/build_sfari_data.py

IMPORTANT — outcome mapping source:
The SFARI document is internally inconsistent. Table 1 (docx table[3], the reference
framework) differs from the mapping used in the doc's own worked example (docx
table[24]) and Appendix B. Only the *example* mapping reproduces the doc's published
sub-indices (Physical 0.55 / Chemical 0.70 / Biological 0.30 / ECI 0.52). We therefore
adopt the example/calculator mapping (parsed from table[24]) as the app default and
record Table 1's mapping alongside it for SME review. Differences: floodplain-connectivity
Physical D->i, hyporheic-connectivity Physical D->i, carbon-processing Chemical D->i.
"""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]                 # sfari/
DATA = ROOT / "data"
DOCX = ROOT / "docs" / "SFARI_Clean.docx"
EASI_FUNCTIONS = Path(r"D:/Code/Work/easi_claude/data/functions.json")

# --- Canonical 20 functions, in document order (ids reused verbatim from EASI). ---
FUNCTIONS = [
    ("catchment-hydrology", "Hydrology", "Catchment hydrology"),
    ("surface-water-storage", "Hydrology", "Surface water storage"),
    ("reach-inflow", "Hydrology", "Reach inflow"),
    ("streamflow-regime", "Hydrology", "Streamflow regime"),
    ("low-flow-baseflow-dynamics", "Hydraulics", "Low flow and baseflow dynamics"),
    ("high-flow-dynamics", "Hydraulics", "High flow dynamics"),
    ("floodplain-connectivity", "Hydraulics", "Floodplain connectivity"),
    ("hyporheic-connectivity", "Hydraulics", "Hyporheic connectivity"),
    ("channel-evolution", "Geomorphology", "Channel evolution"),
    ("channel-floodplain-dynamics", "Geomorphology", "Channel and floodplain dynamics"),
    ("sediment-continuity", "Geomorphology", "Sediment continuity"),
    ("bed-composition-bedform-dynamics", "Geomorphology", "Bed composition and bedform dynamics"),
    ("light-thermal-regime", "Physicochemistry", "Light and thermal regime"),
    ("carbon-processing", "Physicochemistry", "Carbon processing"),
    ("nutrient-cycling", "Physicochemistry", "Nutrient cycling"),
    ("water-soil-quality", "Physicochemistry", "Water and soil quality"),
    ("habitat-provision", "Biology", "Habitat provision"),
    ("population-support", "Biology", "Population support"),
    ("community-dynamics", "Biology", "Community dynamics"),
    ("watershed-connectivity", "Biology", "Watershed connectivity"),
]
FUNCTION_IDS = [f[0] for f in FUNCTIONS]

# docx table indices
TBL_TABLE1 = 3            # Table 1: function -> outcomes (reference framework)
TBL_METRICS = range(4, 24)   # the 20 per-function metric/Likert tables, in order
TBL_NUMCODE = 24         # numerical-code worked example (operative mapping + scores)

LIKERT_ORDER = ["Strongly Agree", "Agree", "Neutral", "Disagree", "Strongly Disagree"]
LIKERT_NUMERIC = {"Strongly Agree": 14, "Agree": 11, "Neutral": 8, "Disagree": 5, "Strongly Disagree": 2}

# Watershed-scale metrics (default is reach 'R').
WATERSHED_SLUGS = {
    "impervious-surface-area", "road-density", "land-use-change", "impoundments",
    "wetland-coverage", "transport-capacity", "fine-sediment-balance",
}

# A few metric tables name a metric slightly differently than its body "Metric N of M:"
# description heading, so the statement won't slug-match. Map table-slug -> heading-slug.
# (Two 5th-metrics — high-flow-velocity-shear-observed, bank-and-floodplain-soil-condition —
# have NO description paragraph in the doc at all and correctly fall back to the metric name.)
METRIC_STATEMENT_ALIASES = {
    "in-channel-ponding-beaver": "in-channel-ponding-and-beaver-activity",
    "flow-permanence-statistics": "flow-statistics",
    "peak-flow-capacity-velocity-shear-stress": "peak-flow-capacity-peak-flow-capacity-velocity-shear-stress",
    "entrenchment-er": "entrenchment-ratio-er",
    "n-p-concentrations": "n-p-concentrations-rapid",
}

# --- Short "reference-good" agreement statements from the paper SFARI Field
# Worksheet v1.0 (docs/FieldForm/Page1-5.jpg), keyed by metricId. These are what
# the assessor agrees/disagrees with in the field; the app shows them next to each
# metric (the longer metricStatement stays in the "how to score" tooltip). Two
# 5th-metrics have no row on the paper form and fall back to metricStatement:
# high-flow-dynamics-high-flow-velocity-shear-observed and
# water-soil-quality-bank-and-floodplain-soil-condition. Transcribed verbatim,
# including en dashes in numeric ranges.
FIELD_STATEMENTS = {
    "catchment-hydrology-impervious-surface-area": "Coverage is minimal, preserving near-natural infiltration/runoff timing, consistent with reference levels.",
    "catchment-hydrology-road-density": "Road density is low enough to avoid significant runoff or sediment inputs, consistent with minimal watershed impact.",
    "catchment-hydrology-land-use-change": "<5% land cover shift in ~15–20 years, indicating stable infiltration/runoff consistent with reference conditions.",
    "catchment-hydrology-impoundments": "Flow is near-natural, with no major dams, unless larger dams are normal per reference.",
    "surface-water-storage-wetland-coverage": "Sufficient wetlands/ponds for flood attenuation/baseflow support, unless minimal wetlands are normal per reference.",
    "surface-water-storage-floodplain-water-retention": "Moderate floods (~1–5 yr) reach the floodplain, providing water retention/infiltration per regional reference.",
    "surface-water-storage-in-channel-ponding-beaver": "Small beaver-type impoundments aid baseflow/habitat; no major fragmentation unless normal per reference.",
    "surface-water-storage-off-channel-storage": "Side channels/beaver ponds/oxbows connect during moderate floods, providing off-channel storage and habitat.",
    "reach-inflow-concentrated-flow-inputs": "Storm/tile drains, ditches are minimal or absent, preserving near-natural infiltration/runoff timing.",
    "reach-inflow-tributary-condition-and-impact": "Tributaries stable, matching reference (flow, quality, form), with no major flash/pollution surges.",
    "reach-inflow-local-runoff-diversions": "Diversions/returns are minimal, flow remains near natural regime without artificial additions/withdrawals.",
    "reach-inflow-road-highway-drainage": "Runoff is minimal or well-managed, causing no significant direct inflows or pollutant pulses.",
    "streamflow-regime-flow-permanence": "Flow duration category from regional SDAM metrics (ephemeral/intermittent/perennial) matches reference. Timing, duration, and frequency of seasonal flows align with reference. No major shift in dryness, baseflow, or mean annual flows, indicating minimal alteration from withdrawals. Streambed vegetation extent and type aligns with expected flow type.",
    "streamflow-regime-flow-permanence-statistics": "Quantitative hydrologic statistics (e.g., flow-duration stats, annual flow days, percent of year flow, flashiness, recession, rate-of-change) fall within reference expectation.",
    "streamflow-regime-channel-natural-flow-regime": "Baseflow to peak floods align with reference condition of magnitude, duration, frequency, and timing of flows indicating minimal hydrologic alteration.",
    "streamflow-regime-artificial-structures-and-inputs": "No major flood-control or flow-regulating structures alter peak/baseflow. Minimal disruption from pumping, dams, or diversion.",
    "low-flow-baseflow-dynamics-low-flow-velocity": "Low-flow velocities, baseflow, and seasonal flow timing match the stream type, reference condition, and historical climate cycles (rain/snowmelt/dry). No unnatural stagnation or dryness compared to reference.",
    "low-flow-baseflow-dynamics-low-flow-depth": "At normal low flow: channel is mostly wetted, exposed channel bed matches the stream's natural wet/dry regime, pools offer sufficient depth for typical fish/invertebrate refuge. Key low/moderate flows align with reference, indicating minimal hydrologic alteration.",
    "low-flow-baseflow-dynamics-longitudinal-connectivity": "At normal low flow, longitudinally: velocity changes, depth changes, movement of water and sediment, thalweg depth, and residual pool connectivity match reference or expected condition.",
    "low-flow-baseflow-dynamics-lateral-connectivity": "At normal low flow, laterally: width of channel and bed wetted, velocity changes, depth changes, movement of water and sediment, thalweg depth and connectivity to floodplain match reference or expected condition.",
    "high-flow-dynamics-overbank-flow-frequency": "Small floods overtop banks approximately every 1–2 years.",
    "high-flow-dynamics-peak-flow-capacity-morphological-check": "The channel cross section accommodates bankfull flows (every ~1–2 years) without destructive events.",
    "high-flow-dynamics-peak-flow-capacity-velocity-shear-stress": "Channel, substrate, and banks accommodate bankfull floods and shear stress without frequent failures or excessive scour/fill, indicating stable capacity.",
    "high-flow-dynamics-bed-mobilization-frequency": "Bed sediments mobilize at expected bankfull flood timing (~1–2 year), maintaining a stable but dynamic channel.",
    "floodplain-connectivity-floodplain-complexity": "Floodplain/off-channel features stay diverse/connected at moderate floods, reflecting minimal loss.",
    "floodplain-connectivity-entrenchment-er": "Channel is not deeply incised; moderate floods access a broad floodplain, matching reference conditions.",
    "floodplain-connectivity-channel-condition": "Stable, not heavily dredged/incised, allowing normal overbank flows/meanders unless dredging is expected per reference.",
    "floodplain-connectivity-lateral-floodplain-inundation": "1–2 yr floods routinely access the floodplain, matching reference inundation patterns.",
    "hyporheic-connectivity-channel-complexity-for-exchange": "Diverse bedforms enhance surface–subsurface exchange, matching reference morphology.",
    "hyporheic-connectivity-floodplain-permeability": "Floodplain soils allow hyporheic/groundwater exchange, consistent with local reference.",
    "hyporheic-connectivity-bed-surface-grain-size": "Substrate has coarse particles with minimal fines, enabling interstitial flow per local references.",
    "hyporheic-connectivity-visible-hyporheic-indicators": "Springs/seeps/subsurface flow show active hyporheic exchange, matching local reference.",
    "channel-evolution-channel-evolution-stage": "Channel is stable/recovered (Stage I/VI or 0,1,6,8), not actively degrading, unless partial evolution is expected per reference.",
    "channel-evolution-incision-trend-headcuts": "No active headcuts or fully stabilized ones; no ongoing incision or bank widening.",
    "channel-evolution-widening-trend": "Bank erosion/migration rates match natural reference; no excessive widening unless expected per reference.",
    "channel-evolution-recovery-indicators": "Channel shows vegetated benches/floodplain formation, indicating recovery from incision/widening.",
    "channel-floodplain-dynamics-bank-erosion-potential": "Banks are under minimal erosion risk (e.g., rate Very Low/Low hazard by BEHI criteria). Only a small fraction of banks erode, unless higher rates are normal in highly dynamic meanders. Natural/gradual erosion, sufficient vegetation, and quick recovery after highs indicate stability. Banks mostly rely on natural deep-rooted vegetation, unless artificial armoring is normal per reference.",
    "channel-floodplain-dynamics-bank-migration-and-meander": "Bank migration and meander shifts occur at a natural rate, unless forced straightening or extreme rates are typical for local geology. No abnormal or human-induced channel jumps.",
    "channel-floodplain-dynamics-sinuosity": "Sinuosity is appropriate for this stream type, unless lower/higher curvature is normal per reference (e.g., steep vs. meandering).",
    "channel-floodplain-dynamics-channel-pattern": "Channel follows a natural planform and remains consistent with reference: minimal straightening, and no unnatural braiding. Secondary/side channels match reference, provide extra habitat/capacity unless blocked historically.",
    "sediment-continuity-sediment-deposition-patterns": "Deposition matches natural form, no severe bar formation beyond typical references.",
    "sediment-continuity-channel-degradation-incision": "Bed is stable, with no active incision disconnecting floodplain unless minor incision is normal.",
    "sediment-continuity-fine-sediment-balance": "Fine input aligns with transport, unless higher turbidity/fines are normal per reference (e.g., blackwater).",
    "sediment-continuity-transport-capacity": "Flow/slope balance incoming sediment, avoiding major bars or deep scour, unless imbalance is normal per reference.",
    "bed-composition-bedform-dynamics-large-wood-frequency-and-diversity": "Large wood frequency is consistent with forested regions, unless minimal per reference (e.g., treeless prairies). Diversity in large wood pieces (rootwads, stable pieces) foster robust habitat in this forested stream.",
    "bed-composition-bedform-dynamics-riparian-wood-recruitment": "The corridor has mature/mid-successional trees ensuring long-term large woody debris input.",
    "bed-composition-bedform-dynamics-substrate-composition": "Classes match reference (no excessive fines or coarsening/armoring beyond typical). Coarse substrate mostly unembedded, unless high fines are normal (wetland streams); embeddedness is healthy.",
    "bed-composition-bedform-dynamics-bedform-diversity": "Riffles, runs, pools (or step–pool), and adequate depths provide habitat complexity matching reference.",
    "light-thermal-regime-riparian-canopy-cover": "Shade is sufficient for this stream type, matching reference condition or unless naturally open.",
    "light-thermal-regime-stream-temperature-rapid": "Stays within acceptable bounds for local aquatic life, rarely exceeding stress thresholds.",
    "light-thermal-regime-algal-growth-light-limited": "Moderate growth matches local light regime, unless high/low is expected per reference (e.g., blackwater).",
    "light-thermal-regime-thermal-refugia": "Springs or deeper pools temper extremes, unless large daily swings are normal per reference (e.g., arid ephemeral).",
    "carbon-processing-cpom-retention": "Leaf litter/coarse organic matter are visibly retained, supporting detrital processes per local references.",
    "carbon-processing-detritus-decomposition-rate-shredder-detritivore-presence": "Visible fungal colonization/shredding indicates a healthy detrital cycle. Quick checks show invertebrates, indicating healthy detrital processing (if typical).",
    "carbon-processing-riparian-corridor-width-and-quality": "Sufficient native vegetation width unless narrow per reference (e.g., prairie streams).",
    "carbon-processing-algal-and-primary-production": "Balanced algae/macrophyte growth, no nuisance blooms, unless moderate coverage is normal per reference.",
    "nutrient-cycling-visible-algal-indications-n-p": "Algae suggests no excessive nutrients vs. reference, unless moderate per reference (e.g., open canopy).",
    "nutrient-cycling-n-p-concentrations": "Nutrients near natural/guidelines unless higher levels are normal per reference (e.g., blackwater).",
    "nutrient-cycling-vegetated-riparian-corridor-width": "Sufficient buffer filters nutrients, matching local references.",
    "nutrient-cycling-relative-denitrification-potential": "Adequate floodplain/wetland pockets for denitrification, consistent with local reference.",
    "water-soil-quality-water-clarity-turbidity": "Appears acceptable, minimal turbidity matching local reference conditions.",
    "water-soil-quality-dissolved-oxygen-rapid": "DO is high enough (>5–6 mg/L) for aquatic life, unless naturally lower in warm/boggy streams.",
    "water-soil-quality-pollutants": "No visible sheen or strong odor, consistent with an unpolluted stream. No chemical/fecal contamination signs; soils appear natural, no suspicious odors/staining.",
    "water-soil-quality-ph-and-specific-conductivity-rapid": "Values near ~6.5–8.5/local norms, no extreme shift unless naturally acidic (bog) or alkaline (arid).",
    "habitat-provision-in-stream-habitat-complexity": "Multiple habitat types (riffles/pools/large wood) exist if expected per reference, not unnaturally uniform.",
    "habitat-provision-overhanging-vegetation": "Provides shelter unless minimal per reference condition in certain ecoregions.",
    "habitat-provision-aquatic-invertebrate-habitat": "Stable substrates (cobble/wood/leaf packs) exist if expected per reference condition (e.g., ephemeral conditions differ). Varied substrates (gravel, cobble, wood) provide multiple habitats, unless uniform is normal.",
    "habitat-provision-lateral-and-off-channel-habitats": "Backwater areas are accessible as nursery/refuge unless naturally confined per reference condition.",
    "population-support-fish-habitat": "Clean gravel/vegetation is present unless alternative spawning is normal per reference condition. Sufficient cover (wood, undercuts, veg) for recruitment/survival if expected per reference condition.",
    "population-support-aquatic-invertebrate-community": "Diverse macroinvertebrates (incl. EPT) if expected per reference, unless ephemeral dryness limits.",
    "population-support-fish-presence": "Multiple native fish species if expected per reference, or unless headwaters were fishless. Juveniles/larvae appear if breeding is normal, unless ephemeral dryness excludes fish.",
    "population-support-amphibian-and-crayfish-presence": "Healthy numbers if expected per reference, unless ephemeral dryness excludes them per reference.",
    "community-dynamics-native-and-non-native-species": "Most historically native fish/inverts remain, unless naturally low-diversity (headwaters, ephemeral). Invasives don't dominate if native-dominated per reference, unless naturalized species cause minimal impact.",
    "community-dynamics-species-richness-composition-and-abundance": "Species richness and composition, trophic composition, fish abundance and composition; or IBI score.",
    "community-dynamics-riparian-communities": "~30 m or appropriate for shade/leaf input if historically expected, unless narrower is normal (e.g., prairie).",
    "community-dynamics-keystone-species": "Species like beaver, unionid mussels, large crayfish persist if expected per reference, unless absent regionally. Invasives don't dominate if historically native-dominated, unless naturalized species cause minimal impact.",
    "watershed-connectivity-upstream-and-downstream-barriers": "No major barriers impede migration unless natural waterfalls exist.",
    "watershed-connectivity-culvert-and-road-crossing-passability": "Crossings allow aquatic organism passage, unless perched/undersized culverts are normal per reference.",
    "watershed-connectivity-lateral-connectivity-for-riparian-fauna": "Floodplain/wetlands accessible for semiaquatic fauna unless steep canyons preclude it.",
    "watershed-connectivity-dewatered-or-intermittent-segments": "No artificial dryness in a historically perennial system, unless ephemeral dryness is normal per reference.",
}

# --- The 26 desktop-supportable metrics (doc Table 22), keyed by metricId. ---
# client "manual" => no clean national source; the evidence adapter returns
# status="unavailable" plus the resource deep-link so the user reviews it by hand.
DESKTOP = {
  "catchment-hydrology-impervious-surface-area": dict(adapter="desktop_hydrology.impervious", client="streamcat", field="pctimp2019ws", label="NLCD 2019 Impervious (MRLC) via Model My Watershed", url="https://modelmywatershed.org/"),
  "catchment-hydrology-road-density": dict(adapter="desktop_hydrology.road_density", client="streamcat", field="rddens", label="Road density — StreamCat rddens / TIGER roads", url="https://www.census.gov/geographies/mapping-files/time-series/geo/tiger-line-file.html"),
  "catchment-hydrology-land-use-change": dict(adapter="desktop_hydrology.land_use_change", client="landcover_change", field="pct_converted", label="NLCD land cover change (2001 vs 2019)", url="https://www.mrlc.gov/viewer/"),
  "catchment-hydrology-impoundments": dict(adapter="desktop_hydrology.impoundments", client="nid_barriers", field="dam_count", label="USACE NID + USGS NABD (upstream dams)", url="https://nid.sec.usace.army.mil/"),
  "surface-water-storage-wetland-coverage": dict(adapter="desktop_hydrology.wetland_coverage", client="streamcat", field="pctwdwet2019ws+pcthbwet2019ws", label="USFWS NWI over FEMA NFHL; StreamCat wetlands", url="https://www.fws.gov/program/national-wetlands-inventory/wetlands-mapper"),
  "reach-inflow-concentrated-flow-inputs": dict(adapter="desktop_hydrology.concentrated_inputs", client="manual", field="", label="Municipal MS4 outfall GIS layers (local)", url="https://www.epa.gov/npdes/stormwater-discharges-municipal-sources"),
  "reach-inflow-local-runoff-diversions": dict(adapter="desktop_hydrology.diversions", client="manual", field="", label="State water-rights / diversions GIS", url="https://waterdata.usgs.gov/"),
  "streamflow-regime-flow-permanence": dict(adapter="desktop_hydrology.flow_permanence", client="nwis", field="zero_flow_days", label="USGS NWIS daily flow (zero-flow days) + NHDPlus HR FCODE", url="https://waterdata.usgs.gov/nwis"),
  "streamflow-regime-flow-permanence-statistics": dict(adapter="desktop_hydrology.flow_statistics", client="nwis", field="flow_duration", label="USGS NWIS daily flow (flow-duration / IHA)", url="https://waterdata.usgs.gov/nwis"),
  "streamflow-regime-channel-natural-flow-regime": dict(adapter="desktop_hydrology.natural_flow_regime", client="nwis", field="flow_alteration", label="TNC IHA on NWIS; else StreamCat dam regulation", url="https://www.conservationgateway.org/"),
  "streamflow-regime-artificial-structures-and-inputs": dict(adapter="desktop_hydrology.artificial_structures", client="aquatic_barriers", field="structures", label="USACE National Levee Database + NID", url="https://levees.sec.usace.army.mil/"),
  "low-flow-baseflow-dynamics-low-flow-depth": dict(adapter="desktop_hydraulics.low_flow_depth", client="xscalc", field="depth", label="3DEP cross-section + Manning (native XS calc) at low-flow Q", url=""),
  "low-flow-baseflow-dynamics-low-flow-velocity": dict(adapter="desktop_hydraulics.low_flow_velocity", client="xscalc", field="velocity", label="Manning (native XS calc): low-flow Q, slope, roughness", url=""),
  "high-flow-dynamics-overbank-flow-frequency": dict(adapter="desktop_hydraulics.overbank_frequency", client="bieger", field="recurrence", label="USGS PeakFQ / regional bankfull-Q regression", url="https://streamstats.usgs.gov/"),
  "high-flow-dynamics-peak-flow-capacity-morphological-check": dict(adapter="desktop_hydraulics.peak_capacity", client="bieger", field="bankfull_geom", label="Regional bankfull-geometry regressions + 3DEP", url="https://streamstats.usgs.gov/"),
  "high-flow-dynamics-bed-mobilization-frequency": dict(adapter="desktop_hydraulics.bed_mobilization", client="xscalc", field="shear", label="Shields (D50 + peak Q) via native XS calc shear", url=""),
  "floodplain-connectivity-lateral-floodplain-inundation": dict(adapter="desktop_hydraulics.lateral_inundation", client="threedep", field="inundation", label="NWI + NAIP + 3DEP DEM/hillshade", url="https://apps.nationalmap.gov/downloader/"),
  "channel-floodplain-dynamics-bank-migration-and-meander": dict(adapter="desktop_geomorph.bank_migration", client="manual", field="", label="Multi-date NAIP/Google Earth; USGS DSAS / RivMAP; EarthExplorer", url="https://earthexplorer.usgs.gov/"),
  "channel-floodplain-dynamics-channel-pattern": dict(adapter="desktop_geomorph.channel_pattern", client="delineation", field="sinuosity", label="Historic USGS topo (TopoView) + NAIP + 3DEP LiDAR; sinuosity", url="https://ngmdb.usgs.gov/topoview/"),
  "sediment-continuity-transport-capacity": dict(adapter="desktop_geomorph.transport_capacity", client="threedep", field="stream_power", label="StreamStats flow + NHDPlus HR slope/length (stream power)", url="https://streamstats.usgs.gov/"),
  "light-thermal-regime-riparian-canopy-cover": dict(adapter="desktop_physicochem.riparian_canopy", client="enviroatlas", field="canopy_pct", label="NLCD Tree Canopy Cover + EPA EnviroAtlas riparian", url="https://enviroatlas.epa.gov/enviroatlas/interactivemap/"),
  "carbon-processing-riparian-corridor-width-and-quality": dict(adapter="desktop_physicochem.riparian_corridor", client="enviroatlas", field="corridor", label="EPA EnviroAtlas riparian (15-50 m) + NLCD", url="https://enviroatlas.epa.gov/enviroatlas/interactivemap/"),
  "nutrient-cycling-vegetated-riparian-corridor-width": dict(adapter="desktop_physicochem.riparian_width", client="enviroatlas", field="width", label="EPA EnviroAtlas riparian", url="https://enviroatlas.epa.gov/enviroatlas/interactivemap/"),
  "community-dynamics-riparian-communities": dict(adapter="desktop_biology.riparian_communities", client="enviroatlas", field="veg", label="EPA EnviroAtlas riparian; StreamCat; NLCD", url="https://enviroatlas.epa.gov/enviroatlas/interactivemap/"),
  "watershed-connectivity-upstream-and-downstream-barriers": dict(adapter="desktop_biology.barriers", client="aquatic_barriers", field="barrier_count", label="National Aquatic Barrier Inventory (USFWS/TNC) + USACE NID", url="https://connectivity.sarpdata.com/"),
  "watershed-connectivity-dewatered-or-intermittent-segments": dict(adapter="desktop_biology.dewatered_segments", client="nwis", field="zero_flow_days", label="USGS NWIS zero-flow days + NHD FCODE + imagery", url="https://waterdata.usgs.gov/nwis"),
}

# Online resources catalog (doc Table 23) — surfaced in the UI / report citations.
ONLINE_RESOURCES = [
    {"name": "EPA EnviroAtlas", "description": "Interactive environmental data incl. stream health, land use, riparian, water quality.", "url": "https://enviroatlas.epa.gov/enviroatlas/interactivemap/", "reference": "U.S. EPA, 2023"},
    {"name": "Google Earth", "description": "Satellite imagery for visualizing streams, watersheds, and land features.", "url": "https://earth.google.com/", "reference": "Google, 2023"},
    {"name": "Model My Watershed", "description": "Watershed modeling; land use, stream and water-quality assessment.", "url": "https://modelmywatershed.org/", "reference": "Stroud Water Research Center, 2023"},
    {"name": "USGS National Map Advanced Viewer", "description": "Topographic and hydrologic data incl. streamflow and watershed boundaries.", "url": "https://apps.nationalmap.gov/viewer/", "reference": "USGS, 2023a"},
    {"name": "USGS StreamStats", "description": "Streamflow statistics and watershed characteristics at a point.", "url": "https://streamstats.usgs.gov/ss/", "reference": "USGS, 2023b"},
]


def slug(s: str) -> str:
    s = (s or "").lower().strip().replace("&", " and ")
    s = re.sub(r"[^a-z0-9]+", "-", s)
    return s.strip("-")


def cell_text(c) -> str:
    return re.sub(r"\s+", " ", c.text).strip()


def parse_metric_table(table, function_id):
    """Group the table's rows into metrics; each metric -> ordered likertCriteria."""
    metrics = []
    current = None
    for row in table.rows[1:]:                       # skip header
        cells = [cell_text(c) for c in row.cells]
        if len(cells) < 4:
            cells = cells + [""] * (4 - len(cells))
        name, observation, criteria, likert = cells[0], cells[1], cells[2], cells[3]
        if not name or name.lower() == "metric":
            # continuation of a merged metric cell without repeated text
            if current is None:
                continue
        elif current is None or name != current["name"]:
            current = {"name": name, "rows": []}
            metrics.append(current)
        current["rows"].append({"likert": likert, "criteria": criteria, "observation": observation})
    return metrics


_METRIC_HEAD_RE = re.compile(r"^\s*Metric\s+\d+\s+of\s+\d+\s*[:.\-–—]\s*(.+?)\s*$", re.I)


def _clean_description(s: str) -> str:
    """Normalize a 'Description. …' paragraph into one clean statement sentence-set:
    drop the leading label and any (Author, 2005; …) literature citations."""
    s = re.sub(r"\s+", " ", s or "").strip()
    s = re.sub(r"^description\s*[.:\-–—]?\s*", "", s, flags=re.I)
    s = re.sub(r"\s*\([^)]*\b\d{4}\b[^)]*\)", "", s)          # (Walsh et al., 2005; …)
    s = re.sub(r"\s+([.,;])", r"\1", s)                        # tidy space left before punctuation
    return re.sub(r"\s{2,}", " ", s).strip()


def parse_metric_descriptions(doc):
    """slug(metric name) -> cleaned metric statement.

    In the docx body each metric is introduced by a 'Metric N of M: NAME' heading
    immediately followed by a 'Description. …' paragraph (these live in the document
    body, not in the Likert tables, so parse_metric_table never sees them)."""
    paras = doc.paragraphs
    out = {}
    for i, p in enumerate(paras):
        head = _METRIC_HEAD_RE.match(re.sub(r"\s+", " ", p.text or ""))
        if not head:
            continue
        name = head.group(1)
        desc = ""
        for q in paras[i + 1:i + 7]:                          # the Description para is the next body line
            qt = re.sub(r"\s+", " ", q.text or "").strip()
            if re.match(r"^description\b", qt, re.I):
                desc = qt
                break
        cleaned = _clean_description(desc)
        if cleaned:
            out[slug(name)] = cleaned
    return out


def build_metrics(doc):
    easi_funcs = {f["id"]: f for f in json.loads(EASI_FUNCTIONS.read_text(encoding="utf-8"))} \
        if EASI_FUNCTIONS.exists() else {}
    descriptions = parse_metric_descriptions(doc)
    out = []
    for fi, tidx in enumerate(TBL_METRICS):
        function_id = FUNCTION_IDS[fi]
        category = FUNCTIONS[fi][1]
        for m in parse_metric_table(doc.tables[tidx], function_id):
            mslug = slug(m["name"])
            metric_id = f"{function_id}-{mslug}"
            likert_criteria = []
            for r in m["rows"]:
                likert_criteria.append({
                    "likert": r["likert"],
                    "criteria": r["criteria"],
                    "exampleObservation": r["observation"],
                })
            binding = DESKTOP.get(metric_id)
            entry = {
                "metricId": metric_id,
                "functionId": function_id,
                "category": category,
                "name": m["name"],
                "scale": "W" if mslug in WATERSHED_SLUGS else "R",
                "metricStatement": (descriptions.get(mslug)
                                    or descriptions.get(METRIC_STATEMENT_ALIASES.get(mslug, ""), "")
                                    or m["name"]),
                "fieldStatement": FIELD_STATEMENTS.get(metric_id, ""),
                "likertCriteria": likert_criteria,
                "desktopSupportable": binding is not None,
            }
            if binding:
                entry["desktopSource"] = {
                    "adapter": binding["adapter"],
                    "client": binding["client"],
                    "field": binding["field"],
                    "label": binding["label"],
                    "url": binding["url"],
                }
                # autoSuggest breaks are curated per-metric in Phase 3; placeholder here.
                entry["autoSuggest"] = None
            out.append(entry)
    return out


def build_outcome_mapping(doc):
    """Operative mapping from the worked-example table (reproduces published sub-indices)."""
    tbl = doc.tables[TBL_NUMCODE]
    mapping = []
    # data rows are r3..r22 (20 functions, in canonical order)
    data_rows = [r for r in tbl.rows[3:3 + 20]]
    for fi, row in enumerate(data_rows):
        cells = [cell_text(c) for c in row.cells]
        p, c, b = cells[3], cells[4], cells[5]
        mapping.append({
            "id": FUNCTION_IDS[fi],
            "physical": p if p in ("D", "i", "-") else "-",
            "chemical": c if c in ("D", "i", "-") else "-",
            "biological": b if b in ("D", "i", "-") else "-",
        })
    return mapping


def build_table1_mapping(doc):
    """Table 1 reference mapping (recorded for SME review; NOT the app default)."""
    tbl = doc.tables[TBL_TABLE1]
    mapping = []
    data_rows = [r for r in tbl.rows[2:2 + 20]]
    for fi, row in enumerate(data_rows):
        cells = [cell_text(c) for c in row.cells]
        p, c, b = cells[2], cells[3], cells[4]
        mapping.append({
            "id": FUNCTION_IDS[fi],
            "physical": p if p in ("D", "i", "-") else "-",
            "chemical": c if c in ("D", "i", "-") else "-",
            "biological": b if b in ("D", "i", "-") else "-",
        })
    return mapping


def build_functions(doc):
    easi = {f["id"]: f for f in json.loads(EASI_FUNCTIONS.read_text(encoding="utf-8"))} \
        if EASI_FUNCTIONS.exists() else {}
    out = []
    for order, (fid, category, name) in enumerate(FUNCTIONS, start=1):
        src = easi.get(fid, {})
        out.append({
            "id": fid,
            "category": category,
            "order": order,
            "name": name,
            "shortDescription": src.get("short_description", ""),
            "definition": src.get("long_description", ""),
            "functionStatement": src.get("function_statement", ""),
            "scoreBand": {"functioning": [11, 15], "atRisk": [6, 10], "nonFunctioning": [0, 5]},
        })
    return out


def main():
    if not DOCX.exists():
        sys.exit(f"docx not found: {DOCX}")
    DATA.mkdir(parents=True, exist_ok=True)
    doc = Document(str(DOCX))

    functions = build_functions(doc)
    mapping = build_outcome_mapping(doc)
    table1 = build_table1_mapping(doc)
    metrics = build_metrics(doc)

    (DATA / "sfari-functions.json").write_text(json.dumps(functions, indent=2, ensure_ascii=False), encoding="utf-8")
    (DATA / "sfari-outcome-mapping.json").write_text(json.dumps(mapping, indent=2, ensure_ascii=False), encoding="utf-8")
    (DATA / "sfari-outcome-mapping-table1.json").write_text(json.dumps(table1, indent=2, ensure_ascii=False), encoding="utf-8")
    metrics_doc = {
        "schemaVersion": 1,
        "method": "Stream Functional Assessment Rapid Index (SFARI)",
        "count": len(metrics),
        "likertScale": {
            "order": LIKERT_ORDER,
            "numeric": LIKERT_NUMERIC,
            "notApplicable": "Not Applicable",
        },
        "onlineResources": ONLINE_RESOURCES,
        "metrics": metrics,
    }
    (DATA / "sfari-metrics.json").write_text(json.dumps(metrics_doc, indent=2, ensure_ascii=False), encoding="utf-8")

    # --- summary ---
    desk = sum(1 for m in metrics if m["desktopSupportable"])
    print(f"functions: {len(functions)}  metrics: {len(metrics)}  desktop-supportable: {desk}")
    with_stmt = [m for m in metrics if (m.get("metricStatement") or "") not in ("", m["name"])]
    print(f"metric statements populated: {len(with_stmt)}/{len(metrics)}")
    for m in metrics:
        if (m.get("metricStatement") or "") in ("", m["name"]):
            print(f"  NO STATEMENT (using name): {m['metricId']}")
    per_fn = {}
    for m in metrics:
        per_fn.setdefault(m["functionId"], 0)
        per_fn[m["functionId"]] += 1
    for fid in FUNCTION_IDS:
        print(f"  {fid}: {per_fn.get(fid, 0)} metrics")
    missing = [k for k in DESKTOP if k not in {m['metricId'] for m in metrics}]
    if missing:
        print("WARNING: desktop bindings with no matching metric:")
        for k in missing:
            print("   ", k)


if __name__ == "__main__":
    main()
